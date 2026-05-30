"""
Unit tests for document_loader.py.
Tests run without a real RAG engine or embedding model.
"""

import sys
from pathlib import Path

# Make project root importable
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest

from src.core.document_loader import (
    DocumentLoadError,
    UnsupportedFileTypeError,
    load_document,
)


# ── TXT ────────────────────────────────────────────────────────────────────

class TestLoadTxt:
    def test_utf8(self, tmp_path):
        f = tmp_path / "utf8.txt"
        f.write_text("Hello 世界", encoding="utf-8")
        text, meta = load_document(f)
        assert "Hello" in text
        assert "世界" in text
        assert meta["file_type"] == "txt"
        assert meta["char_count"] > 0

    def test_gbk_encoding(self, tmp_path):
        f = tmp_path / "gbk.txt"
        f.write_bytes("中文内容".encode("gbk"))
        text, meta = load_document(f)
        assert "中文内容" in text

    def test_latin1_fallback(self, tmp_path):
        f = tmp_path / "latin.txt"
        f.write_bytes(b"\xe9l\xe8ve")  # "élève" in latin-1
        text, meta = load_document(f)
        assert len(text) > 0

    def test_metadata_keys(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("content", encoding="utf-8")
        _, meta = load_document(f)
        for key in ("source_path", "file_name", "file_type", "char_count"):
            assert key in meta


# ── MD ─────────────────────────────────────────────────────────────────────

class TestLoadMd:
    def test_markdown_preserved(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("# Heading\n\n**bold** text", encoding="utf-8")
        text, meta = load_document(f)
        assert "# Heading" in text
        assert "**bold**" in text
        assert meta["file_type"] == "md"


# ── PDF ────────────────────────────────────────────────────────────────────

class TestLoadPdf:
    def test_valid_pdf(self, tmp_path):
        """Create a minimal valid PDF with known text."""
        pytest.importorskip("pypdf")
        import pypdf
        from pypdf import PdfWriter

        writer = PdfWriter()
        page = writer.add_blank_page(width=200, height=200)
        # Add a simple annotation with text so extract_text works
        # For a minimal test, we write a PDF that pypdf can at least open.
        out = tmp_path / "test.pdf"
        with open(out, "wb") as fh:
            writer.write(fh)

        # Should not raise; text may be empty for a blank page
        text, meta = load_document(out)
        assert meta["file_type"] == "pdf"
        assert "page_count" in meta
        assert meta["page_count"] == 1

    def test_corrupted_pdf_raises(self, tmp_path):
        f = tmp_path / "bad.pdf"
        f.write_bytes(b"this is not a PDF")
        with pytest.raises(DocumentLoadError):
            load_document(f)


# ── DOCX ───────────────────────────────────────────────────────────────────

class TestLoadDocx:
    def test_paragraphs(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        out = tmp_path / "test.docx"
        doc.save(str(out))

        text, meta = load_document(out)
        assert "First paragraph." in text
        assert "Second paragraph." in text
        assert meta["file_type"] == "docx"

    def test_table_extraction(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        out = tmp_path / "table.docx"
        doc.save(str(out))

        text, _ = load_document(out)
        for cell in ("A", "B", "C", "D"):
            assert cell in text


# ── Error cases ────────────────────────────────────────────────────────────

class TestErrors:
    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            load_document("/nonexistent/path/file.txt")

    def test_unsupported_extension(self, tmp_path):
        f = tmp_path / "sheet.xlsx"
        f.write_bytes(b"dummy")
        with pytest.raises(UnsupportedFileTypeError):
            load_document(f)
