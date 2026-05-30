"""
Document loaders for PDF, TXT, MD, DOCX, and source code files.
Each loader returns (text: str, metadata: dict).
No chunking happens here — that is chunker.py's responsibility.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Union

from .config import CODE_EXTENSIONS, SUPPORTED_EXTENSIONS

logger = logging.getLogger(__name__)


# ── Custom exceptions ──────────────────────────────────────────────────────

class UnsupportedFileTypeError(ValueError):
    """Raised when the file extension is not in SUPPORTED_EXTENSIONS."""


class DocumentLoadError(RuntimeError):
    """Raised when a supported file cannot be parsed."""


# ── Public API ─────────────────────────────────────────────────────────────

def load_document(file_path: Union[str, Path]) -> tuple[str, dict]:
    """
    Load a document and return its text content plus metadata.

    Returns:
        (text, metadata) where metadata contains:
            source_path, file_name, file_type, char_count,
            and page_count for PDFs.

    Raises:
        FileNotFoundError: if the path does not exist.
        UnsupportedFileTypeError: if the extension is not supported.
        DocumentLoadError: if parsing fails.
    """
    path = Path(file_path).resolve()

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    loaders = {
        ".pdf": _load_pdf,
        ".txt": _load_txt,
        ".md": _load_md,
        ".docx": _load_docx,
    }

    if ext in CODE_EXTENSIONS:
        loader = _load_code
    else:
        loader = loaders.get(ext, _load_txt)

    text, extra_meta = loader(path)

    metadata = {
        "source_path": str(path),
        "file_name": path.name,
        "file_type": ext.lstrip("."),
        "char_count": len(text),
        **extra_meta,
    }

    logger.info("Loaded %s (%d chars)", path.name, len(text))
    return text, metadata


# ── Private loaders ────────────────────────────────────────────────────────

_CHAPTER_RE = re.compile(
    r"^(第[零一二三四五六七八九十百千]+章|Chapter\s+\d+|[A-Z][A-Z\s]{4,})\s*$",
    re.MULTILINE,
)


def _load_pdf(path: Path) -> tuple[str, dict]:
    """Try pdfplumber first (better layout), fall back to pypdf."""
    try:
        import pdfplumber
    except ImportError:
        return _load_pdf_pypdf(path)
    try:
        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)
            pages = [p.extract_text(x_tolerance=2, y_tolerance=2) or "" for p in pdf.pages]
        text = "\n\n".join(p.strip() for p in pages if p.strip())
        text = _CHAPTER_RE.sub(r"\n\n\1", text)
        return text, {"page_count": page_count, "pdf_loader": "pdfplumber"}
    except Exception:
        return _load_pdf_pypdf(path)


def _load_pdf_pypdf(path: Path) -> tuple[str, dict]:
    try:
        import pypdf
    except ImportError as exc:
        raise DocumentLoadError("Neither pdfplumber nor pypdf is installed.") from exc
    try:
        reader = pypdf.PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
        return text, {"page_count": len(reader.pages), "pdf_loader": "pypdf"}
    except Exception as exc:
        raise DocumentLoadError(f"Failed to parse PDF '{path.name}': {exc}") from exc


def _load_txt(path: Path) -> tuple[str, dict]:
    encodings = ["utf-8", "gbk", "latin-1"]
    for enc in encodings:
        try:
            text = path.read_text(encoding=enc)
            logger.debug("Read %s with encoding %s", path.name, enc)
            return text, {"encoding": enc}
        except (UnicodeDecodeError, LookupError):
            continue
    raise DocumentLoadError(
        f"Could not decode '{path.name}' with any of: {encodings}"
    )


def _load_md(path: Path) -> tuple[str, dict]:
    # Markdown is plain text; preserve syntax (embedding model handles it fine).
    text, meta = _load_txt(path)
    return text, meta


def _load_docx(path: Path) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentLoadError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    try:
        doc = Document(str(path))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        text = "\n\n".join(parts)
        # Explicitly release the document to avoid Windows file locks
        del doc
        return text, {}
    except Exception as exc:
        raise DocumentLoadError(f"Failed to parse DOCX '{path.name}': {exc}") from exc


def _load_code(path: Path) -> tuple[str, dict]:
    """
    Load a source code file.
    Prepends a header comment with the file path so retrieved chunks always
    carry their origin, and truncates lines longer than 500 chars (e.g. minified JS).
    """
    raw_text, meta = _load_txt(path)

    lines = raw_text.splitlines()
    line_count = len(lines)

    # Truncate extremely long lines (minified code, base64 blobs, etc.)
    MAX_LINE = 500
    cleaned = []
    for line in lines:
        if len(line) > MAX_LINE:
            cleaned.append(line[:MAX_LINE] + "  # <truncated>")
        else:
            cleaned.append(line)

    # Prepend a file-path header so every chunk knows its source
    header = f"# File: {path.name}\n"
    text = header + "\n".join(cleaned)

    meta["line_count"] = line_count
    meta["language"] = path.suffix.lstrip(".")
    return text, meta
